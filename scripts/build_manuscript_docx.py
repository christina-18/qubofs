"""Build the manuscript as a Word .docx from manuscript_bioinformatics.md.

Output: manuscript_bioinformatics.docx
- US Letter, Arial 11pt, 1-inch margins, double-spaced body (journal-friendly)
- Title, author block with superscripts, affiliations
- Abstract block (Motivation / Results / Availability)
- Sections 1-4 with H1/H2 headings
- Two main tables (hyperparameter grid; AUC summary)
- References numbered
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from pathlib import Path

OUT = Path("/sessions/eager-festive-ptolemy/mnt/outputs/manuscript_bioinformatics.docx")

doc = Document()

# ---------------- Page setup: US Letter, 1-inch margins ----------------
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Default style: Arial 11
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

# Heading styles tweaked
for h, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
    s = doc.styles[h]
    s.font.name = "Arial"
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_para(text="", bold=False, italic=False, size=11, align=None, space_after=0,
             line_spacing=WD_LINE_SPACING.DOUBLE):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p


def add_runs(p, runs):
    """runs: list of (text, dict_of_attrs). attrs: bold, italic, size, super, color, font."""
    for text, attrs in runs:
        r = p.add_run(text)
        r.bold = attrs.get("bold", False)
        r.italic = attrs.get("italic", False)
        if "size" in attrs:
            r.font.size = Pt(attrs["size"])
        if attrs.get("super"):
            r.font.superscript = True
        if attrs.get("color"):
            r.font.color.rgb = RGBColor(*attrs["color"])
        if attrs.get("font"):
            r.font.name = attrs["font"]


# =========================================================================
# Title
# =========================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(12)
title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
tr = title_p.add_run(
    "Cell-type-specific feature selection via Quadratic Unconstrained Binary Optimization "
    "for cross-cohort multiple sclerosis classification from single-cell RNA sequencing"
)
tr.font.size = Pt(15)
tr.bold = True
tr.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# =========================================================================
# Author block
# =========================================================================
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(4)
auth_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
add_runs(auth_p, [
    ("Mizuho Asada", {"bold": True}),
    ("1,2,3", {"super": True, "bold": True}),
    (", ", {}),
    ("Takahisa Mikami", {"bold": True}),
    ("3,4", {"super": True, "bold": True}),
    (", ", {}),
    ("Daisuke Tominaga", {"bold": True}),
    ("5", {"super": True, "bold": True}),
    (", ", {}),
    ("Michael Levy", {"bold": True}),
    ("3,4,*", {"super": True, "bold": True}),
])

# Affiliations (single-spaced)
def add_affil(num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_runs(p, [
        (str(num), {"super": True, "size": 10}),
        (" " + text, {"size": 10}),
    ])

add_affil(1, "Laboratory of Medical Molecular Analytics, Meiji Pharmaceutical University, Tokyo, Japan")
add_affil(2, "Department of Anesthesiology, Institute of Science Tokyo, Tokyo, Japan")
add_affil(3, "Neuroimmunology Clinic and Research Laboratory, Division of Neuroimmunology & Neuroinfectious Disease, Department of Neurology, Massachusetts General Hospital, 65 Landsdowne St., Cambridge, MA 02139, USA")
add_affil(4, "Harvard Medical School, Boston, MA, USA")
add_affil(5, "Division of Mathematical Sciences / Life Informatics, Meiji Pharmaceutical University, Tokyo, Japan")

# Corresponding author + lead
add_para()
p = add_para(line_spacing=WD_LINE_SPACING.SINGLE)
add_runs(p, [
    ("* Corresponding author: ", {"bold": True, "size": 10}),
    ("Michael Levy, MD, PhD — Department of Neurology, Massachusetts General Hospital, "
     "65 Landsdowne St., Lab 500, Cambridge, MA 02139, USA. Email: mlevy11@mgh.harvard.edu",
     {"size": 10}),
])
p = add_para(line_spacing=WD_LINE_SPACING.SINGLE)
add_runs(p, [
    ("Co-corresponding (lead author): ", {"bold": True, "size": 10}),
    ("Mizuho Asada — odakmpha@gmail.com", {"size": 10}),
])
p = add_para(line_spacing=WD_LINE_SPACING.SINGLE, space_after=8)
add_runs(p, [
    ("Dr. Mikami is a Fellow in Neurology at Harvard Medical School and a Fellow in the "
     "BWH-MGH Multiple Sclerosis & Neuroimmunology Fellowship Program at Mass General Brigham.",
     {"italic": True, "size": 10}),
])

# Article type / Subject
p = add_para(line_spacing=WD_LINE_SPACING.SINGLE)
add_runs(p, [("Article type: ", {"bold": True}), ("Original Paper (Methods)", {})])
p = add_para(line_spacing=WD_LINE_SPACING.SINGLE, space_after=12)
add_runs(p, [("Subject section: ", {"bold": True}), ("Genome analysis / Gene expression", {})])

# =========================================================================
# Abstract
# =========================================================================
doc.add_heading("Abstract", level=1)

def abstract_block(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, [(f"{label}: ", {"bold": True}), (text, {})])

abstract_block("Motivation",
    "Single-cell RNA sequencing (scRNA-seq) has accelerated biomarker discovery in multiple sclerosis (MS), "
    "yet cross-cohort reproducibility remains limited by batch effects and donor variability. Existing "
    "feature selection methods optimize either statistical relevance (DE-based selection) or sparsity "
    "(LASSO / Elastic Net) but rarely control redundancy among co-regulated genes within a cell-type-"
    "specific framework. Univariate filters yield panels saturated with co-expressed gene clusters, while "
    "penalized regression confounds redundancy reduction with classifier fitting.")

abstract_block("Results",
    "We introduce a per-cell-type gene panel selection framework based on Quadratic Unconstrained Binary "
    "Optimization (QUBO) that jointly optimizes three complementary axes — relevance, pairwise non-"
    "redundancy, and cardinality — within a single quadratic objective, decoupled from the downstream "
    "classifier. Applied to a four-cohort integrated MS scRNA-seq compendium (50 patients, 99 samples, "
    "385,116 cells), QUBO panels achieve cross-cohort held-out AUC of 0.788 (CSF, σ = 0.044) and 0.768 "
    "(PBMC, σ = 0.033) under Leave-One-Cohort-Out validation, ranking first among five methods (DE-top, "
    "HVG, LASSO, Elastic Net, QUBO). Without prior knowledge, the resulting panels recover all 13 "
    "candidate-pool genes of an independently published B-cell pathogenic signature (Ramesh et al. 2020, "
    "PNAS). Per-cell AUCell scoring of the QUBO B-cell panel produces an MS-versus-HD effect of +0.049 "
    "(q = 5.7×10⁻¹⁵), confirming biology at single-cell resolution.")

abstract_block("Availability and implementation",
    "The pipeline is implemented in Python and R, released under MIT license at "
    "https://github.com/christina-18/scRNA-QUBO, with reproducible Docker / conda environments and "
    "worked examples on the four public cohorts.")

abstract_block("Contact", "odakmpha@gmail.com")
abstract_block("Supplementary information", "Supplementary data are available at Bioinformatics online.")

# =========================================================================
# 1. Introduction
# =========================================================================
doc.add_heading("1. Introduction", level=1)

intro_paras = [
    "Multiple sclerosis (MS) is a chronic autoimmune disease of the central nervous system (CNS) for which "
    "diagnostic and prognostic biomarkers remain limited to MRI imaging and CSF oligoclonal bands. Single-cell "
    "RNA sequencing (scRNA-seq) has accelerated MS biomarker discovery by resolving immune-cell heterogeneity "
    "in cerebrospinal fluid (CSF) and peripheral blood mononuclear cells (PBMC) (Schafflick et al. 2020; "
    "Pappalardo et al. 2020; Ramesh et al. 2020; Heming et al. 2021). However, cross-cohort reproducibility "
    "of candidate gene panels remains poor because of batch effects, sequencing-platform differences, and "
    "donor heterogeneity (Heumos et al. 2023). A reproducible, cell-type-specific biomarker framework is "
    "required for translation to clinical use.",

    "Existing feature-selection strategies fall into two camps. Univariate filters — e.g. taking the top-K "
    "differentially expressed genes (DE-top) by edgeR, DESeq2, or limma-voom |t-statistic|, or top-K highly "
    "variable genes (HVG) — are simple but prone to selecting clusters of co-expressed genes (e.g. multiple "
    "HLA-class-II family members) that contribute redundant information. Penalized regression (LASSO, "
    "Elastic Net) sparsifies via L1 / L1+L2 regularization but ties redundancy reduction to a specific "
    "classifier loss, conflating selection with model fitting; selection is also unstable across resamples "
    "(Meinshausen & Bühlmann 2010). Neither family enforces a hard cardinality constraint with explicit "
    "pairwise dissimilarity, both of which are essential for designing biomarker panels that are "
    "interpretable, clinically tractable, and reproducible.",

    "Cheminformatics provides a complementary template: in compound library design, relevance × diversity × "
    "cardinality is jointly optimized (Snarey et al. 1997; Pearlman & Smith 1998). The natural mathematical "
    "home for this triplet is Quadratic Unconstrained Binary Optimization (QUBO) — a binary-variable "
    "quadratic minimization framework with mature classical (Simulated Annealing, Tabu Search) and quantum "
    "(D-Wave) solvers (Lucas 2014; Glover et al. 2018).",

    "Here we adapt the QUBO framework to scRNA-seq biomarker selection. Our contributions are:",
]
for t in intro_paras:
    add_para(t)

bullets = [
    "We formulate per-cell-type biomarker selection as a QUBO with three terms — score-weighted relevance, correlation-penalized redundancy, and soft cardinality — and a hyperparameter grid auto-selected by inner cross-validation.",
    "We integrate QUBO selection into a soft-voting cell-type ensemble classifier, evaluated under Leave-One-Cohort-Out (LOCO) cross-validation across four public MS scRNA-seq cohorts (50 patients, 385,116 cells).",
    "We benchmark QUBO against four selection baselines (DE-top, HVG, LASSO, Elastic Net) under matched candidate pool, K grid, classifier, and ensemble — isolating the selection-logic effect.",
    "We provide independent biological validation: recovery of all candidate-pool genes from a published B-cell signature (Ramesh et al. 2020), and cell-level AUCell scoring confirming MS-versus-HD effect (q = 5.7×10⁻¹⁵).",
    "We disclose a methodological limitation — pseudobulk dilution of T-cell signal — and outline a Multi-Instance Learning extension as future work.",
]
for i, b in enumerate(bullets, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(b)

add_para("The pipeline is open-source at https://github.com/christina-18/scRNA-QUBO.")

# =========================================================================
# 2. Materials and Methods
# =========================================================================
doc.add_heading("2. Materials and Methods", level=1)

doc.add_heading("2.1 Datasets and pre-processing", level=2)
add_para(
    "We integrated four publicly available MS scRNA-seq cohorts: Pappalardo et al. 2020 (PRJNA671484, 11 "
    "donors, CSF + PBMC), Heming et al. 2021 (osmzhlab/MS-ence-cov, 18 donors, CSF longitudinal), Ramesh "
    "et al. 2020 (PRJNA549712, 17 donors, CSF + PBMC), and Touil et al. 2023 (PRJNA979258, 4 donors, CSF "
    "cryopreserved, healthy donors only). Total 50 unique donors, 99 sample-tissue combinations, 385,116 "
    "cells. Cohorts span 22 healthy donors (HD; 64% female, age 34.1 ± 8.3) and 28 MS patients (64% female, "
    "age 37.2 ± 11.0); demographics are matched across diagnosis (Supplementary Table S1)."
)
add_para(
    "Raw count matrices were obtained from each archive, harmonized to gene symbols (HGNC nomenclature), "
    "and integrated as a single Seurat object (Seurat v5.0). Standard quality control (mitochondrial "
    "fraction < 20%, ≥ 200 genes/cell) and SCTransform normalization were applied per cohort prior to "
    "integration."
)

doc.add_heading("2.2 Cell-type annotation, biology filter, and pseudobulk", level=2)
add_para(
    "Cells were annotated using Azimuth predicted.celltype.l2 (Hao et al. 2021) and collapsed into eight "
    "broad immune subsets: B, Mono, CD4_T, CD8_T, NK, DC, dnT, gdT. Three criteria justified this "
    "resolution: (1) major MS-relevant lymphoid / myeloid populations; (2) ≥ 20 cells per donor on average "
    "for stable pseudobulk aggregation; (3) annotation reproducibility across cohorts."
)
add_para(
    "Genes were restricted to biology-relevant features by removing mitochondrial (MT-), ribosomal "
    "(RPL/RPS except RPLP0/1/2 and RPSA — laminin-receptor / blood-brain-barrier transmigration components "
    "retained for MS relevance), heat-shock (HSP), nuclear lncRNAs (MALAT1, NEAT1), and a curated "
    "housekeeping list, following Heumos et al. 2023 best practices. Approximately 7,960 genes remained."
)
add_para(
    "For each (donor × cell type × tissue) triplet, mean pseudobulk expression vectors were constructed. "
    "The donor became the unit of statistical analysis. Pseudobulk profiles were further filtered to the "
    "top 3,000 highly variable features per cell type (Seurat FindVariableFeatures, vst method)."
)

doc.add_heading("2.3 QUBO formulation", level=2)
add_para(
    "For each (cell type × tissue × LOCO fold), we constructed a QUBO with binary decision variables "
    "x_i ∈ {0, 1} indicating whether candidate gene i is selected. The cost function is:"
)
# Equation as bold centered text
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
r = p.add_run("H(x) = − sᵀ x + γ · xᵀ R x + λ · (Σᵢ xᵢ − K)²")
r.font.size = Pt(13)
r.bold = True

add_para("where:")
qubo_terms = [
    "s ∈ ℝᴺ: per-gene relevance score computed as the squared edgeR |t-statistic| for the MS-vs-HD contrast (with covariates: log10(n_cells), age, sex, batch). The candidate pool N = 100 comprises the top-100 genes by |t|.",
    "R ∈ ℝᴺˣᴺ: pairwise redundancy matrix, computed as |Pearson correlation| across pseudobulk donor profiles (off-diagonal entries only). Self-similarity is excluded.",
    "γ: redundancy penalty weight (γ ∈ {1.0} in our reported results, evaluated up to γ ∈ {0.5, 1.0, 2.0} in sensitivity analysis).",
    "λ: cardinality penalty enforcing Σᵢ xᵢ ≈ K.",
    "K ∈ {10, 20, 30}: target cardinality, auto-selected per (cell type × fold) by inner 5-fold cross-validation on the training cohorts (see §2.4).",
]
for b in qubo_terms:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(b)

add_para(
    "Expanding the cardinality term yields a standard QUBO upper-triangular matrix Q, which is passed "
    "unchanged to a generic solver. Importantly, the same matrix is solvable on classical and "
    "quantum-annealing hardware, ensuring portability."
)

doc.add_heading("2.4 Solver and hyperparameter selection", level=2)
add_para(
    "We solve the QUBO with classical Simulated Annealing (dwave-neal) using 100 reads × 1,000 sweeps per "
    "instance (~3 s on a single core; total runtime 8–12 min for the full 5-method × 8 cell-type × "
    "2-tissue × 3-fold benchmark). Quantum annealing is left as an option for future scale-up but is not "
    "required."
)
add_para(
    "For each (cell type × outer LOCO fold), inner 5-fold CV on the outer training donors selects "
    "hyperparameters by maximizing held-out inner-fold AUC. The chosen hyperparameters are then re-fit on "
    "the full outer training data and applied to the held-out cohort."
)

# Table 2: Hyperparameter grids
add_para()
p = add_para("Table 2. Hyperparameter grids and selection-determinism per method.",
             bold=True, size=10, line_spacing=WD_LINE_SPACING.SINGLE)
t2 = doc.add_table(rows=6, cols=3)
t2.style = "Light Grid Accent 1"
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ["Method", "Tuned by inner 5-fold CV", "Selection determinism given training data"]
for i, h in enumerate(hdr):
    cell = t2.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)

t2_rows = [
    ["DE-top", "K ∈ {10, 20, 30}", "Deterministic — gene ranking by |edgeR t| is unique"],
    ["HVG", "K ∈ {10, 20, 30}", "Deterministic — gene ranking by training-set variance is unique"],
    ["LASSO", "K ∈ {10, 20, 30} × C grid (5 values)", "Stochastic in C — selected genes vary with regularization path"],
    ["Elastic Net", "K × C grid × l1_ratio = 0.5", "Stochastic in C"],
    ["QUBO", "K ∈ {10, 20, 30} × γ ∈ {1.0} (sensitivity {0.5, 1.0, 2.0})", "Stochastic in SA — multiple-run consensus available"],
]
for ri, row in enumerate(t2_rows, 1):
    for ci, val in enumerate(row):
        c = t2.rows[ri].cells[ci]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)

add_para()
add_para(
    "For DE-top and HVG, gene combinations are deterministic given the outer training data; inner CV "
    "serves solely to select K. For LASSO/Elastic Net and QUBO, inner CV jointly tunes K and the "
    "additional method-specific knobs. All five methods consume the same matched candidate pool and feed "
    "the same downstream classifier and ensemble (§2.5), so the comparison isolates the effect of the "
    "selection logic, not the cardinality budget or fitting algorithm."
)
add_para(
    "To stabilize QUBO selection across the stochastic SA, a consensus mode was also evaluated: 10 "
    "independent SA runs with seed offsets, returning the top-K genes by selection frequency. Both raw "
    "and consensus QUBO yielded equivalent cross-cohort AUC; raw QUBO is reported in the main text."
)

doc.add_heading("2.5 Per-cell-type ensemble classifier and evaluation", level=2)
add_para(
    "For each cell type, an L2-regularized logistic regression (scikit-learn, C = 1.0) was fit on the "
    "QUBO-selected panel (mean-pseudobulk inputs), producing donor-level MS probability p_c for cell type "
    "c. Per-patient probabilities were combined by soft voting (simple mean across cell types where "
    "≥ 1 gene was selected)."
)
add_para(
    "External evaluation used Leave-One-Cohort-Out (LOCO) cross-validation across the 3 MS-containing "
    "cohorts (Pappalardo, Heming, Ramesh), yielding 3 hold-out folds per tissue. Touil (HD only) was "
    "fixed to the training set throughout. PBMC LOCO included only Pappalardo and Ramesh (Heming has no "
    "PBMC samples), yielding 2 hold-out folds."
)
add_para(
    "Reported metrics: AUC (ROC), AP (precision-recall), F1, MCC, and σ_AUC (standard deviation across "
    "folds — a direct measure of cross-cohort stability)."
)

doc.add_heading("2.6 Baseline methods", level=2)
add_para(
    "Four baselines were evaluated under matched conditions (same candidate pool, same K grid, same "
    "classifier, same ensemble): DE-top (top-K by |edgeR t|, deterministic); HVG (top-K by training-set "
    "variance, no MS/HD label, deterministic); LASSO (L1-penalized logistic, C tuned to yield ~K non-zero "
    "coefficients then top-K by |coefficient|); Elastic Net (L1 + L2, l1_ratio = 0.5, top-K by "
    "|coefficient|). Forcing all methods to the same K grid and downstream classifier/ensemble isolates "
    "the selection-logic difference from cardinality and fitting-algorithm differences."
)

doc.add_heading("2.7 Biological validation", level=2)
add_para("Three orthogonal validations were performed:")
val_items = [
    "(i) Curated gene-set enrichment: Hypergeometric tests of QUBO-selected gene unions against five MS-relevant gene sets (Hametner 2013 iron rim, Reactome MHC II, Ramesh 2020 B-cell, IMSGC 2019 GWAS, van Langelaar 2020 Type I IFN), with FDR-BH correction.",
    "(ii) Independent literature signature recovery (non-circular): comparison of QUBO panels against the 27-gene B-cell signature reported by Ramesh et al. 2020 (PNAS), which was not used in our pipeline. The candidate pool retained 13 of the 27 (the remaining 14 were filtered by HVG / biology / low expression).",
    "(iii) Cell-level AUCell validation (Aibar et al. 2017): per-cell scoring of seven literature-curated MS gene sets with Wilcoxon BH-FDR significance testing. We additionally scored the QUBO-selected panels themselves at single-cell resolution.",
]
for v in val_items:
    add_para(v)

# =========================================================================
# 3. Results
# =========================================================================
doc.add_heading("3. Results", level=1)

doc.add_heading("3.1 Cross-cohort biomarker performance", level=2)
add_para(
    "Across the four-cohort LOCO design, QUBO achieved the highest held-out AUC in both compartments: "
    "CSF AUC = 0.788 (σ = 0.044), PBMC AUC = 0.768 (σ = 0.033) (Table 1). The cross-cohort spread σ_AUC "
    "for QUBO (0.044) was markedly tighter than LASSO (0.068, +55%), comparable to Elastic Net (0.041), "
    "and considerably better than DE-top (0.065) and HVG (0.048). On three additional metrics (F1, MCC, "
    "AP), QUBO ranked first on AUC, F1, and MCC — losing AP narrowly to Elastic Net (Δ = 0.024, n.s.). "
    "This indicates that the QUBO advantage is driven by improved feature combinations — not by superior "
    "cardinality or classifier choice, since these were matched."
)
add_para(
    "Per-cohort breakdown (CSF): QUBO scored Pappalardo 0.807, Heming 0.738, Ramesh 0.819 — all within "
    "0.74–0.82, contrasted with LASSO's wider 0.72–0.85 spread. The tight range supports clinical "
    "translatability, where cross-site stability is essential."
)

# Table 1: AUC summary
add_para()
p = add_para("Table 1. Cross-cohort held-out AUC and σ_AUC summary by method × tissue (CSF).",
             bold=True, size=10, line_spacing=WD_LINE_SPACING.SINGLE)
t1 = doc.add_table(rows=6, cols=5)
t1.style = "Light Grid Accent 1"
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ["Method", "AUC mean", "σ (cohort)", "Genes/ct/fold", "Prediction panel (5 ct/fold)"]
for i, h in enumerate(hdr):
    c = t1.rows[0].cells[i]
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)

t1_rows = [
    ["DE-top", "0.742", "0.065", "15.5", "77"],
    ["HVG", "0.712", "0.048", "16.4", "82"],
    ["LASSO", "0.779", "0.068", "10.7", "53"],
    ["Elastic Net", "0.779", "0.041", "16.9", "85"],
    ["QUBO (proposed)", "0.788", "0.044", "14.7", "73"],
]
for ri, row in enumerate(t1_rows, 1):
    for ci, val in enumerate(row):
        c = t1.rows[ri].cells[ci]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
        if ri == 5:
            r.bold = True

add_para()

doc.add_heading("3.2 Cell-type-specific gene panels recover MS biology axes", level=2)
add_para(
    "QUBO selected on average 17 genes per cell type per fold (range 10–30). The union across all panels "
    "was 448 unique genes; 21 genes formed a stable core recurrently selected in ≥ 50% of per-cell-type "
    "panels."
)
add_para(
    "Heatmap visualization (Figure 2) of the top-5 union panels per cell type across 5 effective cell "
    "types (B, Mono, NK, dnT, gdT) shows that each panel recovers its native cell-type biology:"
)
biology = [
    "B: plasma / IgM (IGHM) + secretory machinery (SPCS2);",
    "Mono: iron-rim biology (FTL, FTH1) + MHC II (HLA-DPB1, IFI30, CD74) + myeloid defense (CST3, LYZ);",
    "NK: cytotoxic axis (KLRB1, KLRC1, CCL5, CRIP1, LTB);",
    "dnT: cytotoxic + Type I IFN (GZMA, ISG15, IL32);",
    "gdT: cytotoxic + activation (TPT1, SRGN, CD69, GZMA).",
]
for b in biology:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(b)
add_para(
    "Curated gene-set enrichment of the stable core showed iron metabolism (Hametner 2013) at fold "
    "enrichment 36× (q = 2×10⁻³), cytotoxic effectors at 18×, and MHC II pathway at 16× — all significant."
)

doc.add_heading("3.3 External validation against published signatures", level=2)
add_para(
    "We compared QUBO selections against the 27-gene clonally-expanded pathogenic B-cell signature "
    "reported by Ramesh et al. 2020 (PNAS), which was not used at any stage of our pipeline. Of the 13 "
    "Ramesh genes retained in our candidate pool (the remaining 14 were excluded at the biology / HVG / "
    "expression-level filter), all 13 were present in our QUBO panels (hypergeometric test: fold "
    "enrichment 2.44, p = 8.3×10⁻⁶). A second curated signature (CSF immune dynamics, 22 genes) yielded "
    "9/12 recovery (fold enrichment 1.83, p = 0.018)."
)
add_para(
    "This convergence with two methodologically independent published signatures provides non-circular "
    "validation that QUBO selections capture biologically meaningful axes — not merely statistical "
    "artifacts."
)

doc.add_heading("3.4 Per-cell AUCell scoring confirms biology at single-cell resolution", level=2)
add_para(
    "Independent of QUBO selection, we scored seven literature-curated MS gene sets across all eight cell "
    "types using AUCell (Aibar et al. 2017) at single-cell resolution. Expected MS pathology axes were "
    "recovered cell-type-specifically: CD8 T × Type I IFN (q = 1.3×10⁻⁷); NK × cytotoxic (q < 10⁻⁵); "
    "Mono × MHC II + iron rim (q < 10⁻⁵)."
)
add_para(
    "We additionally scored the QUBO-selected panels themselves per cell. The B-cell QUBO panel produced "
    "an MS-vs-HD median-activity difference of +0.049 with q = 5.7×10⁻¹⁵, the strongest single-cell-level "
    "effect across all cell types — direct evidence that the panel discriminates MS at single-cell "
    "granularity, not merely at the donor pseudobulk level. Mono and NK panels showed similarly "
    "significant cell-level effects (q < 10⁻⁸)."
)

doc.add_heading("3.5 Limitation: pseudobulk dilution in T cells", level=2)
add_para(
    "In CSF, CD4_T, CD8_T, and DC produced zero QUBO-selected genes because the candidate pool itself was "
    "exhausted by the HVG / biology / DE filters (CD4_T retained 1 candidate gene, CD8_T retained 0, DC "
    "retained 0). This affected all five methods equally — it is not a QUBO failure. The mechanism is "
    "pseudobulk dilution: CD4 and CD8 are mixtures of functional subtypes (Th1/Th17/Treg/Tfh; "
    "Tem/Tcm/Trm/exhausted/MAIT), and donor-level pseudobulk averaging dilutes minority disease-driving "
    "subset signals (Th17, CD8 effectors) within the ~90% baseline T-cell population. CD4_T was the most "
    "abundant CSF cell type (98,000 cells) yet contributed nothing; this defines the most important "
    "methodological limitation of the present approach."
)

# =========================================================================
# 4. Discussion
# =========================================================================
doc.add_heading("4. Discussion", level=1)
add_para(
    "We have introduced QUBO-based per-cell-type biomarker selection for scRNA-seq, the first framework "
    "that jointly optimizes relevance, pairwise non-redundancy, and cardinality within a single "
    "quadratic objective decoupled from the downstream classifier. Across four MS cohorts (50 patients, "
    "385K cells, LOCO design), QUBO achieved the highest cross-cohort AUC and the second-tightest σ_AUC "
    "among five matched methods, while recovering known MS biology axes (iron rim, MHC II, cytotoxic, "
    "Type I IFN, plasma cell). External validation by recovery of an independently published 13-gene "
    "B-cell signature (all 13 candidate-pool genes recovered) and single-cell AUCell scoring of QUBO "
    "panels (q = 5.7×10⁻¹⁵ for B-cell discrimination) provide non-circular biological evidence that the "
    "panels are interpretable and meaningful — not just statistical optima."
)
p = add_para()
add_runs(p, [("Comparison to penalized regression. ", {"bold": True}),
             ("LASSO and Elastic Net implicitly handle redundancy via L1/L2 regularization, but redundancy "
              "reduction is coupled to classifier loss, making selection unstable across resamples and "
              "dependent on the regularization path. QUBO decouples selection from fitting, allowing "
              "arbitrary downstream classifiers and explicit, interpretable redundancy control via the "
              "gene-pair correlation matrix.", {})])
p = add_para()
add_runs(p, [("Comparison to univariate filters. ", {"bold": True}),
             ("DE-top and HVG produce panels saturated with co-expressed clusters (e.g. multiple "
              "HLA-class-II family members ranking jointly), wasting cardinality budget. QUBO's "
              "γxᵀRx penalty actively spreads selection across distinct pathways, evidenced by the "
              "cell-type-specific recovery of multiple MS axes per panel.", {})])
p = add_para()
add_runs(p, [("Limitations. ", {"bold": True}),
             ("First, pseudobulk dilution of T-cell signal is fundamental and affects all donor-level "
              "methods, not QUBO specifically — but it is the most consequential limit in the CSF "
              "compartment. Second, simulated annealing is heuristic, and runtime grows quadratically "
              "with candidate-pool size; we mitigate by capping pool to top-100 by |t|, which suffices "
              "for our setting but may need adaptation for genome-wide selection. Third, EDSS, disease "
              "duration, and DMT history were unavailable from public metadata; clinical "
              "phenotype-stratified extension is left to Phase 2 with site collaborations.", {})])
p = add_para()
add_runs(p, [("Future directions. ", {"bold": True}),
             ("(1) Multi-Instance Learning (MIL): each donor as a bag of cells, attention-weighted to "
              "surface disease-driving subsets (Th17, Trm) without pseudobulk averaging — directly "
              "resolving the T-cell dilution. QUBO retains gene selection and gains a new role in "
              "informative-cell coreset selection per donor. (2) Cohort expansion: integration of "
              "recently published cohorts (Jacobs et al. 2024 Cell Reports Medicine, 354K CSF cells, 123 "
              "untreated MS; Ban et al. 2024 Brain, 97K CSF cells with eQTL annotations) under 6-fold "
              "LOCO. (3) Quantum solver scale-up: D-Wave or Hybrid-CQM solvers for genome-wide QUBO "
              "without candidate-pool truncation.", {})])

# =========================================================================
# Funding / COI / Acknowledgements / Data Availability
# =========================================================================
doc.add_heading("Funding", level=1)
add_para("This work received no specific funding from any agency in the public, commercial, or not-for-profit sectors.")

doc.add_heading("Conflicts of Interest", level=1)
add_para("The authors declare no competing interests.")

doc.add_heading("Acknowledgements", level=1)
add_para(
    "The authors thank Pappalardo, Heming, Ramesh, and Touil and their respective groups for making "
    "their scRNA-seq data publicly available, which made the cross-cohort integration in this study "
    "possible. We also thank the BWH-MGH Multiple Sclerosis & Neuroimmunology Fellowship Program at "
    "Mass General Brigham for providing the research environment in which this work was carried out."
)

doc.add_heading("Data Availability", level=1)
add_para(
    "All four input cohorts are available at the original repositories: Pappalardo et al. 2020 "
    "(PRJNA671484); Heming et al. 2021 (osmzhlab repository); Ramesh et al. 2020 (PRJNA549712); Touil "
    "et al. 2023 (PRJNA979258). Pre-processed pseudobulk matrices, all selected gene lists, intermediate "
    "results, and full reproducibility scripts are released at https://github.com/christina-18/scRNA-QUBO "
    "with version-pinned conda and Docker environments."
)

# =========================================================================
# References
# =========================================================================
doc.add_heading("References", level=1)
refs = [
    "Aibar, S. et al. (2017) SCENIC: single-cell regulatory network inference and clustering. Nat. Methods, 14, 1083–1086.",
    "Ban, M., Bredikhin, D., Huang, Y. et al. (2024) Expression profiling of cerebrospinal fluid identifies dysregulated antiviral mechanisms in multiple sclerosis. Brain, 147, 554–565.",
    "Glover, F., Kochenberger, G., Du, Y. (2018) A Tutorial on Formulating and Using QUBO Models. arXiv:1811.11538.",
    "Hametner, S. et al. (2013) Iron and neurodegeneration in the multiple sclerosis brain. Ann. Neurol., 74, 848–861.",
    "Hao, Y. et al. (2021) Integrated analysis of multimodal single-cell data. Cell, 184, 3573–3587.",
    "Heming, M. et al. (2021) Neurological manifestations of COVID-19 feature T cell exhaustion and dedifferentiated monocytes in CSF. Immunity, 54, 164–175.",
    "Heumos, L. et al. (2023) Best practices for single-cell analysis across modalities. Nat. Rev. Genet., 24, 550–572.",
    "Jacobs, B.M. et al. (2024) Single-cell analysis of cerebrospinal fluid reveals common features of neuroinflammation. Cell Rep. Med., 6, 101733.",
    "Lucas, A. (2014) Ising formulations of many NP problems. Front. Phys., 2, 5.",
    "Meinshausen, N., Bühlmann, P. (2010) Stability selection. J. R. Stat. Soc. B, 72, 417–473.",
    "Pappalardo, J.L. et al. (2020) Transcriptomic and clonal characterization of T cells in the human CNS. Sci. Immunol., 5, eabb8786.",
    "Ramesh, A. et al. (2020) A pathogenic and clonally expanded B cell transcriptome in active multiple sclerosis. PNAS, 117, 22932–22943.",
    "Schafflick, D. et al. (2020) Integrated single cell analysis of blood and cerebrospinal fluid leukocytes in multiple sclerosis. Nat. Commun., 11, 247.",
    "Snarey, M. et al. (1997) Comparison of algorithms for dissimilarity-based compound selection. J. Mol. Graph. Model., 15, 372–385.",
    "Touil, T. et al. (2023) [Cryopreserved CSF reference dataset]. (Citation TBD)",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    add_runs(p, [(f"{i}. ", {"bold": True, "size": 10}), (r, {"size": 10})])

# =========================================================================
# Figure & Table list
# =========================================================================
doc.add_heading("Figure & Table List", level=1)
items = [
    ("Figure 1", "Cross-cohort held-out AUC by method (CSF / PBMC), with σ_AUC error bars."),
    ("Figure 2", "Heatmap of QUBO-selected gene panels — selection frequency × cell type, with MS pathology axis annotation."),
    ("Figure 3", "Conceptual pipeline overview (per-cell-type QUBO + soft voting)."),
    ("Figure 4", "Per-cell AUCell heatmap of curated MS gene sets across 8 cell types."),
    ("Table 1", "Cross-cohort AUC and σ_AUC summary by method × tissue."),
    ("Table 2", "Hyperparameter grids and selection-determinism per method."),
]
for label, desc in items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(4)
    add_runs(p, [(f"{label}. ", {"bold": True, "size": 10}), (desc, {"size": 10})])

add_para()
add_para("Supplementary:", bold=True, line_spacing=WD_LINE_SPACING.SINGLE)
supp = [
    ("Table S1", "Patient demographics by cohort."),
    ("Table S2", "Per-cell-type top selected genes (frequency-ranked)."),
    ("Table S3", "Hypergeometric enrichment results (full)."),
    ("Table S4", "Gene-count comparison across the 5 selection methods (CSF) — per-fold and per-cohort union."),
    ("Figure S1", "Per-cohort AUC distributions by method."),
    ("Figure S2", "QUBO formulation details and SA convergence."),
    ("Figure S3", "Cell-type coverage diagnosis (CD4/CD8 dilution)."),
    ("Figure S4", "External validation against Ramesh 2020 signature."),
    ("Figure S5", "Per-cell AUCell of QUBO panels."),
]
for label, desc in supp:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(4)
    add_runs(p, [(f"{label}. ", {"bold": True, "size": 10}), (desc, {"size": 10})])

# Save
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
